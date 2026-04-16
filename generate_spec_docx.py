#!/usr/bin/env python3
"""Generate AI_Smart_Inbox_Spec.docx from the restructured spec."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ---- Page setup ----
for section in doc.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# ---- Style setup ----
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    if level == 1:
        hs.font.size = Pt(22)
        hs.paragraph_format.space_before = Pt(24)
        hs.paragraph_format.space_after = Pt(8)
    elif level == 2:
        hs.font.size = Pt(16)
        hs.paragraph_format.space_before = Pt(18)
        hs.paragraph_format.space_after = Pt(6)
    else:
        hs.font.size = Pt(13)
        hs.paragraph_format.space_before = Pt(14)
        hs.paragraph_format.space_after = Pt(4)

# ---- Helpers ----
def tbl(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '1A1A2E')
        shd.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shd)
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, rd in enumerate(rows):
        for ci, ct in enumerate(rd):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ct
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
                run.font.name = 'Calibri'
            if ri % 2 == 1:
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), 'F5F5FA')
                shd.set(qn('w:val'), 'clear')
                cell._tc.get_or_add_tcPr().append(shd)
    doc.add_paragraph()

def quote(text, who):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f'\u201c{text}\u201d')
    r.italic = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.5)
    p2.paragraph_format.space_after = Pt(10)
    r2 = p2.add_run(f'- {who}')
    r2.font.size = Pt(10); r2.font.color.rgb = RGBColor(0x66, 0x66, 0x66); r2.italic = True

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix); r.bold = True; r.font.size = Pt(11); r.font.name = 'Calibri'
        r = p.add_run(f' {text}'); r.font.size = Pt(11); r.font.name = 'Calibri'
    else:
        r = p.add_run(text); r.font.size = Pt(11); r.font.name = 'Calibri'

def body(text):
    doc.add_paragraph(text)

def bold_body(bold_part, rest):
    p = doc.add_paragraph()
    r = p.add_run(bold_part); r.bold = True
    r = p.add_run(rest)

def centered_bold(text, size=11, color=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color

# ============================================================
# TITLE
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(120)
r = p.add_run('AI Smart Inbox'); r.font.size = Pt(32); r.bold = True; r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('for Invoice Collection'); r.font.size = Pt(20); r.font.color.rgb = RGBColor(0x44, 0x44, 0x66)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
r = p.add_run('Product Specification'); r.font.size = Pt(14); r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(30)
r = p.add_run('Brian Nunes De Souza  |  April 2026'); r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ============================================================
# 1. THE PROBLEM
# ============================================================
doc.add_heading('1. The Problem', level=1)

body('After a home service professional sends an invoice, the conversation doesn\'t end. It often just begins. Customers ask questions ("Why was I charged for floor leveling?"), make promises ("I\'ll pay Friday"), express concerns ("Can I split this up?"), or simply go quiet. These messages land in the same inbox as scheduling requests, job updates, and general chatter. For a busy pro who\'s on a job site all day, the payment-critical messages get buried. Every day of silence costs money.')
body('This is not a payments infrastructure problem. HouseCall Pro already supports cards, ACH, financing, BNPL, and tap-to-pay. The tools to pay exist. What\'s missing is the conversation that gets customers to actually pay.')

doc.add_heading('Customer Discovery', level=2)
body('I started this project where good product work should always start: by talking to real customers. My father and stepmother run Rauen Flooring, a small flooring company in Fort Myers, FL with just three employees. I sat down with them and walked through their entire invoice-to-payment workflow, from estimate through final collection.')

quote("I'll finish a job covered in dust, check my phone, and there are six messages. The one about the invoice? Buried.", "Sandro")
quote("Customers always ask 'why did you charge for leveling?' We've explained it a hundred times, but when we're on a job we can't stop and type it out.", "Fabiana")
quote("I thought Sandro answered her. He thought I did. She waited three days and hired someone else.", "Fabiana")

body('Five findings shaped this product:')
bullet('Because flooring requires knowing house dimensions for pricing, invoices are often well-defined and sent before or soon after work begins. The "question window" opens early, and customers review charges before work is even complete.', 'Invoice timing is front-loaded in flooring.')
bullet('When customers receive an invoice, they frequently ask questions, many of which were already discussed in earlier conversations or are common FAQs. "Why did you charge me for float?" is routine for the pro but novel and confusing for the customer. They delay paying until they get clarity.', 'Customer questions are the #1 bottleneck to payment.')
bullet('When they\'re in the middle of a flooring installation, they can\'t stop to answer invoice questions on their phone. By the time they\'re done for the day, the customer\'s message is buried under other notifications.', 'Busy pros forget to follow up.')
bullet('Even in a 3-person company, messages fall through the cracks because everyone assumes someone else responded. In larger companies with dedicated office staff, this compounds: the person who answers messages can\'t answer billing questions and has to escalate to the owner, creating additional delay.', 'The "someone else will handle it" problem.')
bullet('This was the most critical insight. My family was clear: as customers themselves, they don\'t like interacting with AI. They value the personal relationship with their service provider. "I want to know I\'m talking to a person, not a robot." This fundamentally shaped the approach: AI assists the pro behind the scenes. It never speaks to the customer directly.', 'Pros do not want AI talking directly to their customers.')

doc.add_heading('Industry Validation', level=2)
body('After grounding the problem in firsthand research, I investigated whether these patterns hold across other home service verticals (plumbing, HVAC, electrical, cleaning, general contracting). The data strongly confirmed what I heard at the kitchen table:')
bullet('of small businesses are owed money from unpaid invoices, averaging $17,500 per business (QuickBooks 2025 Late Payments Report).', '56%')
bullet('94% at 30 days past due, 74% at 90 days, and only 50% at 6 months (Commercial Collection Agency Association).', 'Collection probability drops sharply with time:')
bullet('(Xero). AI-generated reminders get paid 5 days faster (QuickBooks 2025).', 'One reminder increases payment likelihood by ~50%')
bullet('about delinquent bills for fear of damaging the relationship (Old National Bank, 2025).', '60% of small business owners avoid confronting customers')
bullet('is the #1 reason for non-payment in home services (SCORE/SBA).', 'Dispute over work scope or pricing')
body('Speed and clarity are the two biggest levers for getting paid.')

doc.add_heading('Where the Funnel Leaks', level=2)
body('Think of the invoice lifecycle as a funnel:')
centered_bold('Invoice Sent  \u2192  Customer Receives  \u2192  Questions / Concerns  \u2192  Pro Responds  \u2192  Payment Made')
body('Today, the funnel leaks at the "Questions / Concerns" stage. Customers delay paying until their questions are answered. Pros delay answering because they\'re busy. The longer the gap, the lower the collection probability.')
body('Competitors see this gap too. ServiceTitan now offers AI invoice summaries and one-click collection emails. Jobber has AI rewrite for invoice messages and two-way texting. Workiz provides AI messaging across channels tailored to customer history. But these are point solutions: each one automates a single touchpoint. None have built a unified AI layer that watches the full conversation, understands context from persistent memory, and helps the pro respond at the moment it matters most. That is the opportunity.')

doc.add_heading('What Good Looks Like', level=3)
bullet('Pro opens their inbox and immediately sees the 2-3 most important payment-related actions. No hunting.')
bullet('Each action has an AI-drafted message ready to send, grounded in actual conversation context and invoice details.')
bullet('Pro can see WHY the AI recommended this action (source messages, reasoning chain, confidence level).')
bullet('Customers get timely, helpful responses within hours, not days.')
bullet('Average days-to-payment decreases. The aging invoice tail shortens.')

doc.add_heading('What Failure Looks Like', level=3)
bullet('AI surfaces irrelevant or low-confidence actions. Noise instead of signal.')
bullet('Drafted messages contain wrong amounts, wrong customer details, or tone-deaf language.')
bullet('Pro loses trust and ignores the AI entirely.')
bullet('Customers receive responses that feel robotic or impersonal.')
bullet('Dismiss rates are high, indicating recommendations aren\'t valuable.')

# ============================================================
# 2. THE SOLUTION
# ============================================================
doc.add_heading('2. The Solution', level=1)

body('AI Smart Inbox adds an intelligent layer on top of HouseCall Pro\'s messaging inbox. It surfaces high-priority, payment-related actions as cards at the top of the inbox. Each card represents an AI-identified opportunity to move an invoice toward payment.')

doc.add_heading('What the Pro Experiences', level=2)
bold_body('Action Cards ', 'appear above the conversation list. Horizontally scrollable, each card represents a specific AI recommendation with a priority level (high, medium, low) and a one-tap path to the AI-drafted response.')
bold_body('In-Conversation Drafts ', 'appear at the bottom of the chat thread when a pro opens a conversation that has an AI recommendation. The draft sits in context alongside the actual conversation, so the pro can verify it makes sense before sending.')
bold_body('Reasoning Transparency ', 'is available for every recommendation. A "Why this?" button opens a panel showing the specific customer messages the AI analyzed, the reasoning chain, the business knowledge and FAQ entries that informed the draft, and a confidence score.')

doc.add_heading('The Four Action Types', level=2)
bold_body('Payment Follow-up. ', 'Customer said "I\'ll pay by Friday" and Friday has passed. AI drafts a friendly follow-up referencing the customer\'s own words about when they\'d pay. Supports granular timing: "end of day," "this Friday," "next week."')
bold_body('Invoice Question Response. ', 'Customer asked about a specific charge (e.g., "What\'s the float charge?"). AI drafts an explanation using the actual invoice line-item details, business FAQ, and job context. Only surfaces when AI confidence is high enough to give an accurate answer.')
bold_body('Invoice Summary. ', 'When a new invoice is sent, AI drafts a customer-friendly breakdown highlighting key line items, the total, and how to pay. Warmer and more informative than the raw invoice notification. This is proactive: it fires without a customer message.')
bold_body('Payment Plan Suggestion. ', 'Customer expressed financial difficulty ("I can\'t pay this all at once"). AI drafts a response acknowledging their concern and presenting available options like installment plans or financing through existing HCP integrations.')

doc.add_heading('AI Does vs. Human Does', level=2)
tbl(['AI Does', 'Human Does'], [
    ['Scans conversations for payment-related signals', 'Decides whether to send, edit, or dismiss'],
    ['Identifies follow-up timing based on customer promises', 'Makes judgment calls on complex disputes'],
    ['Drafts contextual, personalized responses', 'Approves every outbound message'],
    ['Surfaces invoice details relevant to customer questions', 'Negotiates payment plans and custom terms'],
    ['Ranks actions by priority and confidence', 'Owns the customer relationship'],
])
body('AI drafts. The pro decides. No message is ever sent without explicit pro approval. This is non-negotiable in the MVP and comes directly from customer feedback.')

doc.add_heading('Why AI, Not Rules', level=2)
body('A rule-based system could handle some of these cases with regex patterns and templates. But customer messages are natural language: "I\'ll get you that money Friday," "paying end of day," and "can I settle up next week?" all mean the same thing and require natural language understanding to parse reliably. Writing a good response requires understanding the specific charge being questioned, the customer\'s tone, and the invoice details. Connecting "What\'s the float charge?" to the correct invoice line item ("Floor leveling/float, $450") is inherently a semantic task. And a rule-based system would need hundreds of hand-crafted rules while still missing edge cases. AI generalizes from context.')

# ============================================================
# 3. HOW IT WORKS
# ============================================================
doc.add_heading('3. How It Works', level=1)

doc.add_heading('Design Principles', level=2)
body('Three principles shaped every backend decision:')
bold_body('Grounded in data. ', 'Every AI output references specific invoice amounts, message IDs, and FAQ entries. No facts are generated from the model\'s training data alone. If the AI references "$450 for floor leveling," the reasoning panel shows exactly which invoice line item it came from.')
bold_body('Structured outputs. ', 'The AI returns typed JSON via tool_use, not free-form text. The application can always parse the result reliably, and every field has a defined schema.')
bold_body('Graceful degradation. ', 'If the AI is unavailable, keyword-based fallbacks keep the inbox functional. Template responses are less intelligent but factually correct. No blank screens or errors are ever shown to the pro.')

doc.add_heading('The Two-Stage Pipeline', level=2)
body('When a customer sends a message, the system runs a two-stage pipeline. Each stage is a separate AI call with its own system prompt, memory context, and structured output schema.')
centered_bold('Customer Message \u2192 Evaluate (Stage 1) \u2192 Action Card? \u2192 Pro Taps "AI Draft" \u2192 Generate (Stage 2) \u2192 Pro Reviews & Sends', size=10, color=RGBColor(0x00, 0x7A, 0xFF))

bold_body('Stage 1: Evaluate. ', 'The system sends the full conversation, invoice details, pro memory, and customer memory to the AI. The AI decides: Does this message warrant an action? What type? What priority? Is there enough context to generate a confident draft? Should the customer\'s memory profile be updated with new information from this message? If the answer is yes, an action card appears in the inbox. If no, the message is treated as routine.')
bold_body('Stage 2: Generate. ', 'When the pro taps "AI Draft," the system selects a specialized prompt based on the action type and calls the AI with the full context. The AI returns the draft message text (2-5 sentences, conversational tone), the reasoning behind the recommendation, the specific source messages it analyzed, the business knowledge entries that informed the draft, and a confidence score between 0.0 and 1.0.')
body('The two-stage design is deliberate. Evaluation runs on every incoming message and must be fast and conservative. Draft generation only runs when the pro explicitly asks for it, so it can take more time and use richer context.')

doc.add_heading('The Memory System', level=2)
body('Memory is what makes this system agentic rather than stateless. Instead of treating every message in isolation, the AI builds and maintains its own understanding of each business and each customer relationship over time.')
bold_body('Pro Memory ', 'is shared across all conversations. It contains the business\'s team members and roles, service area, pricing table, standard policies, and FAQ entries. The FAQ section is generated by the AI from real customer questions: when multiple customers ask about floor leveling charges, the system creates an FAQ entry that future drafts can reference. The pro can review and edit this knowledge base at any time.')
bold_body('Customer Memory ', 'is per-customer and builds incrementally with every message. During each evaluation, the AI decides whether the customer\'s memory should be updated. A message like "I\'m over on Pine Island Road" updates the customer\'s location. "You did my kitchen last year" updates the relationship history. Over time, each customer profile becomes a rich context document that the AI uses to personalize every draft.')
body('This is why a draft responding to Mike\'s question about the float charge can reference the specific conversation where Sandro mentioned the subfloor issue during the estimate visit, cite the FAQ entry explaining what floor leveling is, and note that $450 for 480 sq ft falls below the standard $2-$4/sq ft range. The AI has the full picture because memory gave it the full picture.')

doc.add_heading('Skills Architecture', level=2)
body('The backend defines five distinct AI skills, each implemented as a system prompt paired with a forced structured output schema:')
tbl(['Skill', 'When It Runs', 'What It Does'], [
    ['Evaluate Message', 'Customer sends a new message', 'Decides whether to create an action card, and what type'],
    ['Evaluate Conversation', 'Seed pipeline processes a conversation holistically', 'Same as above, but considers the entire thread state'],
    ['Generate Draft', 'Pro taps "AI Draft"', 'Produces the draft text, reasoning, sources, and confidence'],
    ['Generate Customer Memory', 'Pipeline initializes a customer profile', 'Creates a full memory document from conversation history'],
    ['Generate FAQ', 'After all messages are processed', 'Derives FAQ entries from actual customer questions'],
])
body('Each skill is self-contained: a system prompt and a structured output schema. The pipeline, memory system, and UI all work without modification when a new skill is added.')

doc.add_heading('Provenance and Transparency', level=2)
body('Every AI output is traceable. When a pro taps "Why this?" on a draft, they see the complete evidence chain:')
bullet('the exact customer messages that triggered the action, with timestamps', 'Source messages:')
bullet('the specific FAQ entries or business facts that informed the draft (e.g., "FAQ: What is a float charge?" or "Services & Pricing")', 'Memory sources:')
bullet('1-2 sentences explaining why the AI chose this approach', 'Reasoning chain:')
bullet('displayed as a visual bar, so the pro can gauge certainty before deciding to send', 'Confidence score:')
body('Nothing is a black box. The pro can verify the AI "did its homework" before any message reaches a customer.')

# ============================================================
# 4. TRUST AND AUTONOMY
# ============================================================
doc.add_heading('4. Trust and Autonomy', level=1)

doc.add_heading('The Core Principle', level=2)
body('When AI isn\'t confident, it stays silent. Surfacing a bad recommendation erodes trust faster than surfacing nothing. Below-threshold actions are never shown to the pro.')
body('This principle came directly from customer research. My family told me they\'d stop using a feature the moment it embarrassed them with a customer. One wrong amount, one tone-deaf follow-up to a sensitive situation, and the feature is dead to them. The trust bar in home services is extremely high because these are real, ongoing relationships built on years of service.')

doc.add_heading('Earning Autonomy Through Results', level=2)
body('The MVP is Phase 1: every AI draft requires human review and approval. This is deliberate. Trust must be built through demonstrated accuracy and usefulness before AI earns more responsibility.')
tbl(['Phase', 'AI Behavior', 'Pro Experience'], [
    ['Phase 1 (This MVP)', 'All drafts require manual review and explicit send', 'Pro learns what AI can do, builds confidence through daily use'],
    ['Phase 2 (Earned)', 'Pro can enable auto-send for specific, proven action types', 'AI proved itself on simpler tasks first'],
    ['Phase 3 (Future)', 'Fully autonomous follow-ups for routine cases, with pro notification', 'AI operates independently where trust is established'],
])
body('Phase 2 is not on a fixed timeline. It requires 90+ days of sustained quality metrics at scale, explicit pro demand for more automation, and zero customer complaints about AI quality. Each action type graduates independently based on its own track record. Invoice summaries might auto-send first because they\'re low-risk and informational. Payment follow-ups might stay human-reviewed longer because the stakes are higher.')

doc.add_heading('Guardrails', level=2)
bold_body('Full visibility. ', 'The drafted message text is displayed in an editable text area. Nothing is hidden or sent automatically.')
bold_body('Reasoning transparency. ', 'A "Why this?" button shows the AI\'s logic and the specific source messages it analyzed.')
bold_body('Edit always available. ', 'The pro can modify any part of the draft before sending. The AI\'s suggestion is a starting point, not a final answer.')
bold_body('Factual grounding. ', 'All amounts, dates, and line items are sourced from structured invoice data, never generated from the model\'s training data alone. 100% accuracy on financial details is a hard requirement. Zero tolerance for errors. A single wrong dollar amount could damage a pro\'s credibility and create a billing dispute.')
bold_body('Dismiss with signal. ', 'A dismiss button removes the recommendation. Dismissals are tracked at two levels: individual card dismissals ("this specific suggestion isn\'t helpful") and section-level dismissals ("I don\'t want AI suggestions at all"). If per-card dismiss rate exceeds 20%, we tighten confidence thresholds. If section dismissals exceed 5%, we pause rollout and investigate.')

# ============================================================
# 5. MEASUREMENT
# ============================================================
doc.add_heading('5. Measurement', level=1)

doc.add_heading('Outcome Metrics', level=2)
body('These measure whether the product moves the business needle:')
tbl(['Metric', 'Definition', 'Target'], [
    ['Days-to-payment', 'Avg days from invoice sent to payment received', 'Decrease by 20%+'],
    ['Response rate', '% of customer payment messages with a pro response within 24h', 'Increase from ~40% to 80%+'],
    ['30-day collection rate', '% of invoices paid within 30 days of sending', 'Increase by 15%+'],
    ['Revenue recovered', 'Dollar amount of invoices paid after AI-assisted follow-up', 'Track absolute $'],
    ['Pro time saved', 'Minutes/week pro spends on payment-related messaging', 'Decrease by 50%+'],
])

doc.add_heading('AI Quality Metrics', level=2)
body('These measure whether the AI itself is doing a good job:')
tbl(['Metric', 'Definition', 'Quality Bar'], [
    ['Draft send rate', '% of AI drafts sent (as-is or edited) vs. dismissed', '> 70%'],
    ['Draft acceptance rate', '% of AI drafts sent without any editing', '> 50%'],
    ['Edit rate', '% of drafts edited before sending', '20-40% (healthy range)'],
    ['Dismissal rate', '% of action cards dismissed without sending', '< 20%'],
    ['Reasoning view rate', '% of actions where pro clicks "Why this?"', 'Track (expect higher early)'],
    ['Factual accuracy', '% of drafts with correct invoice amounts, dates, line items', '100% (hard requirement)'],
])

doc.add_heading('Guardrail Metrics', level=2)
body('If any guardrail is breached, we take immediate action rather than waiting for the next review cycle:')
tbl(['Metric', 'Red Line', 'Action if Breached'], [
    ['Per-card dismiss rate', '> 20%', 'Investigate quality; tighten confidence threshold'],
    ['Section dismiss rate', '> 5%', 'Pause rollout; conduct qualitative research'],
    ['Customer CSAT', 'Drop > 5% vs. control', 'Pause feature; investigate conversation quality'],
    ['Factual error rate', 'Any occurrence', 'Immediate hotfix; root cause analysis; incident review'],
    ['Pro NPS on AI feature', '< 30', 'Re-evaluate approach before expanding'],
])

doc.add_heading('Quality Bar Before Shipping', level=2)
body('We do not proceed to wider rollout until all of the following are met in the A/B experiment:')
bullet('Draft send rate > 70%')
bullet('Dismissal rate < 20%')
bullet('Zero factual errors over a 30-day window')
bullet('Pro NPS on AI feature > 30')
bullet('No customer complaints about AI-generated message quality or tone')
bullet('Statistically significant improvement in days-to-payment in treatment vs. control')

# ============================================================
# 6. REQUIREMENTS
# ============================================================
doc.add_heading('6. Requirements', level=1)

doc.add_heading('P0: Must Have (MVP)', level=2)
tbl(['ID', 'Requirement', 'Rationale'], [
    ['FR-01', 'AI scans conversations to identify payment-related signals (promises, questions, concerns, new invoices)', 'Core value proposition: surface what matters without pro effort'],
    ['FR-02', 'Action cards displayed at top of inbox for high-confidence recommendations with priority ranking', 'Immediate visibility; pro sees actions without searching'],
    ['FR-03', 'AI generates contextual draft responses using conversation history and structured invoice data', 'Core feature: ready-to-send messages grounded in real data'],
    ['FR-04', 'Pro can review, edit, and send AI-drafted messages', 'Human-in-the-loop control; no autonomous sending in Phase 1'],
    ['FR-05', 'Reasoning transparency: source messages, confidence score, and logic visible via "Why this?" panel', 'Trust-building through transparency; pro verifies AI\'s work'],
    ['FR-06', 'Dismiss functionality for individual action cards', 'Pro control; dismissals tracked as quality signal'],
    ['FR-07', 'Factual grounding: all amounts, dates, and line items sourced from structured invoice data', 'Financial accuracy is non-negotiable'],
])

doc.add_heading('P1: Should Have', level=2)
tbl(['ID', 'Requirement', 'Rationale'], [
    ['FR-08', 'Four distinct action types: payment follow-up, invoice question, invoice summary, payment plan', 'Different situations need different response strategies and tone'],
    ['FR-09', 'Priority ranking of action cards (high / medium / low) with visual differentiation', 'Help pros triage; highest-impact actions are most visible'],
    ['FR-10', 'Section-level dismiss (opt out of all AI suggestions)', 'Respect pro autonomy; strong negative signal for guardrail monitoring'],
    ['FR-11', 'Dismiss event tracking for analytics and guardrail metrics', 'Quantitative signal for whether recommendations provide value'],
])

doc.add_heading('P2: Nice to Have (Post-MVP)', level=2)
tbl(['ID', 'Requirement', 'Rationale'], [
    ['FR-12', 'Thumbs up / down feedback on individual drafts', 'Explicit quality signal to complement implicit send/dismiss data'],
    ['FR-13', 'Draft improvement learning from pro edit patterns', 'Continuous quality improvement based on how pros modify suggestions'],
    ['FR-14', 'Customizable confidence threshold per pro', 'Power users may want more or fewer suggestions'],
    ['FR-15', 'Auto-send for specific action types with pro opt-in (Phase 2)', 'Earned autonomy for proven, low-risk action types'],
])

doc.add_heading('Non-Functional Requirements', level=2)
tbl(['ID', 'Requirement', 'Target'], [
    ['NFR-01', 'Draft generation latency', '< 5 seconds from request to display'],
    ['NFR-02', 'Action card refresh frequency', 'Within 1 minute of new incoming messages'],
    ['NFR-03', 'Platform support', 'Mobile-first (iOS and Android); responsive web'],
    ['NFR-04', 'Accessibility', 'WCAG 2.1 AA compliance'],
    ['NFR-05', 'Data privacy', 'No customer PII used in AI model training; all processing via API'],
    ['NFR-06', 'Availability', '99.9% uptime; graceful degradation if AI service is unavailable'],
])

# ============================================================
# 7. ROLLOUT
# ============================================================
doc.add_heading('7. Rollout', level=1)

doc.add_heading('Phase 0: Internal Dogfooding (2 weeks)', level=2)
body('HCP product and engineering team tests with synthetic conversations. Identify obvious UX issues, edge cases, and prompt failures before any customer exposure. Build internal confidence and gather team feedback on draft quality.')

doc.add_heading('Phase 1a: Closed Beta (4-6 weeks)', level=2)
body('50-100 hand-selected pros across verticals (plumbing, HVAC, electrical, cleaning, flooring). Mix of solo operators and small teams to test both persona types. Weekly 15-minute feedback sessions with participating pros.')
bold_body('Gate to proceed: ', 'draft send rate > 60%, dismiss rate < 25%, zero factual errors, qualitative feedback positive (NPS > 20).')

doc.add_heading('Phase 1b: A/B Experiment (6-8 weeks)', level=2)
body('~1,000 pros randomly assigned to treatment (AI Smart Inbox) vs. control (standard inbox). Measure days-to-payment, response rate, collection rate, and pro satisfaction. Statistical significance required on primary metrics before proceeding.')
bold_body('Gate to proceed: ', 'draft send rate > 70%, dismiss rate < 20%, zero factual errors, NPS > 30, statistically significant improvement in days-to-payment.')

doc.add_heading('Phase 1c: General Availability', level=2)
body('Gradual rollout: 10% \u2192 25% \u2192 50% \u2192 100%. Feature flag for instant kill switch if guardrails are breached. Continuous monitoring of all guardrail metrics at each stage. Qualitative feedback collection through in-app surveys at 7 and 30 days post-enablement.')

doc.add_heading('Phase 2: Earned Autonomy (Future)', level=2)
body('Phase 2 is not on a fixed timeline. Prerequisites: 90+ days of sustained quality metrics at scale, explicit pro demand for more automation, zero customer complaints about AI quality. Scope: enable opt-in auto-send for specific, low-risk action types (e.g., invoice summaries). Each action type graduates independently. Pros can opt in or out at any time; autonomy is always revocable.')

doc.add_heading('Phase Transition Summary', level=2)
tbl(['Gate', 'Criteria to Pass'], [
    ['Internal \u2192 Beta', 'No critical bugs; team consensus on draft quality'],
    ['Beta \u2192 A/B', 'Send rate > 60%; dismiss < 25%; zero factual errors; NPS > 20'],
    ['A/B \u2192 GA', 'Significant improvement in days-to-payment; send rate > 70%; dismiss < 20%; NPS > 30'],
    ['Phase 1 \u2192 Phase 2', '90+ days sustained quality; explicit pro demand; zero customer complaints'],
])

# ============================================================
# 8. THE BIGGER PICTURE
# ============================================================
doc.add_heading('8. The Bigger Picture', level=1)

body('This MVP is small on purpose. The point is to show the thinking and the foundation, not to ship everything at once. But the architecture behind it is designed for something much larger.')

doc.add_heading('A Platform, Not a Feature', level=2)
body('HouseCall Pro already has incredibly rich data about every customer relationship: conversation history, invoice details, job records, customer profiles. Today, the pro has to mentally synthesize all of this to decide what to do and what to say. The AI Smart Inbox shows what happens when you give an AI layer access to that context: it can identify what matters, draft a relevant response, and have it ready before the pro even opens the app.')
body('The messaging inbox is the first entry point. But the same context model works anywhere a customer interaction needs a fast, accurate, contextual response. Emails. Notifications. Scheduling workflows. Estimate follow-ups. Every new entry point adds signal, and richer context enables more capable AI behaviors.')

doc.add_heading('Skills Multiply', level=2)
body('The four action types in this MVP are the first skills. Each one is a system prompt paired with a structured output schema. Adding a new skill means writing a new prompt, not rebuilding infrastructure.')
body('The next skills are already obvious from customer research: estimate follow-ups, scheduling nudges, repeat customer recognition, warranty question responses. Beyond that, cross-conversation pattern detection could surface insights across the entire customer base ("three customers this week asked about floor leveling charges, you might want to update your estimate template").')

doc.add_heading('Memory Compounds', level=2)
body('The AI gets better at understanding each customer and each business over time. A pro who has been using the system for six months has a richer, more accurate AI assistant than one who started yesterday. Customer memory builds with every message. Business knowledge grows as the FAQ learns from real customer questions. This compounding effect is a durable advantage that deepens with usage.')

doc.add_heading('Beyond Invoicing', level=2)
body('The two-stage pipeline (evaluate, then generate) is not specific to invoices. It works for any workflow where a customer message needs a contextual, accurate response from a busy professional. The same architecture could power AI assistance across scheduling, estimates, job follow-ups, and any other HCP workflow where speed and context matter.')
body('This prototype proves the concept. The architecture is ready for what comes next.')

# ============================================================
# SAVE
# ============================================================
output = 'docs/process/AI_Smart_Inbox_Spec.docx'
doc.save(output)
print(f'Saved {output}')
