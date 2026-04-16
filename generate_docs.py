#!/usr/bin/env python3
"""Generate process documents for HouseCall Pro AI Smart Inbox MVP."""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = "/Users/briandesouza/Desktop/housecall-pro-ai-mvp/process"
AUTHOR = "Brian Nunes De Souza"
DATE = "April 2026"
FONT_NAME = "Calibri"


# ── Helpers ──────────────────────────────────────────────────────────────────

def new_doc():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x1C, 0x1C, 0x1E)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.space_before = Pt(2)
    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = FONT_NAME
        hs.font.color.rgb = RGBColor(0x00, 0x61, 0xFF)
    return doc


def title_page(doc, title, subtitle=""):
    for _ in range(6):
        doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(26)
    r.font.name = FONT_NAME
    r.font.color.rgb = RGBColor(0x00, 0x61, 0xFF)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(subtitle)
        r2.font.size = Pt(16)
        r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        r2.font.name = FONT_NAME
    doc.add_paragraph("")
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rm = meta.add_run(f"{AUTHOR}\n{DATE}")
    rm.font.size = Pt(12)
    rm.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    rm.font.name = FONT_NAME
    doc.add_page_break()


def para(doc, text):
    return doc.add_paragraph(text)


def bold_para(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    return p


def mixed(doc, parts):
    """parts: list of (text, bold) tuples."""
    p = doc.add_paragraph()
    for text, bold in parts:
        r = p.add_run(text)
        r.bold = bold
    return p


def bullets(doc, items):
    for item in items:
        if isinstance(item, tuple):
            # (bold_prefix, rest)
            p = doc.add_paragraph(style="List Bullet")
            r1 = p.add_run(item[0])
            r1.bold = True
            p.add_run(item[1])
        else:
            doc.add_paragraph(item, style="List Bullet")


def numbered(doc, items):
    for item in items:
        if isinstance(item, tuple):
            p = doc.add_paragraph(style="List Number")
            r1 = p.add_run(item[0])
            r1.bold = True
            p.add_run(item[1])
        else:
            doc.add_paragraph(item, style="List Number")


def table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    # Header row
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.name = FONT_NAME
        # Light blue background
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "E8F0FE")
        cell.paragraphs[0].runs[0].element.getparent().getparent().find(
            qn("w:tcPr")
        ) if cell._tc.find(qn("w:tcPr")) is not None else None
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_pr.append(shading)
    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.name = FONT_NAME
    doc.add_paragraph("")  # spacer
    return t


# ── Document 1: Process Document ────────────────────────────────────────────

def create_process_document():
    doc = new_doc()
    title_page(doc, "AI Smart Inbox\nfor Invoice Collection", "Process Document")

    # ── 1. PROBLEM FRAMING ──────────────────────────────────────────────────
    doc.add_heading("1. Problem Framing", level=1)

    doc.add_heading("The Pain Point", level=2)
    para(doc,
        "After a home service professional sends an invoice, the conversation doesn't end "
        "\u2014 it often just begins. Customers ask questions (\"Why was I charged for floor "
        "leveling?\"), make promises (\"I'll pay Friday\"), express concerns (\"Can I split "
        "this up?\"), or simply go quiet. These messages land in the same inbox as scheduling "
        "requests, job updates, and general chatter. For a busy pro who's on a job site all "
        "day, the payment-critical messages get buried \u2014 and every day of silence costs money."
    )
    para(doc,
        "This is not a payments infrastructure problem. HouseCall Pro already supports cards, "
        "ACH, financing, BNPL, and tap-to-pay. The tools to pay exist. What's missing is the "
        "conversation that gets customers to actually pay."
    )

    doc.add_heading("How I Got Here: Customer Discovery", level=2)
    para(doc,
        "I started this project where good product work should always start \u2014 by talking "
        "to real customers. My father and stepmother run a small flooring company with just "
        "three employees. I sat down with them and walked through their entire invoice-to-payment "
        "workflow, from estimate through final collection. Their experiences became the foundation "
        "for this product concept."
    )

    bold_para(doc, "Key findings from these interviews:")
    bullets(doc, [
        ("Invoice timing is front-loaded in flooring. ",
         "Because flooring requires knowing house dimensions for pricing, invoices are often well-defined "
         "and sent before or soon after work begins. This means the \"question window\" opens early "
         "\u2014 customers review charges before work is even complete."),
        ("Customer questions are the #1 bottleneck to payment. ",
         "When customers receive an invoice, they frequently ask questions \u2014 many of which were "
         "already discussed in earlier conversations or are common FAQs. For example: \"Why did you "
         "charge me for float?\" (a leveling compound applied to uneven subfloors). These questions "
         "are routine for the pro but novel and confusing for the customer. Customers delay paying "
         "until they get clarity."),
        ("Busy pros forget to follow up. ",
         "My family described how, when they're in the middle of a flooring installation, they can't "
         "stop to answer invoice questions on their phone. By the time they're done for the day, the "
         "customer's message is buried under other notifications. They lose the customer's momentum "
         "to pay."),
        ("The \"someone else will handle it\" problem. ",
         "Even in their 3-person company, messages fall through the cracks because everyone assumes "
         "someone else responded. In larger companies with dedicated office staff, this compounds \u2014 "
         "the person who answers messages often can't answer billing questions and has to escalate to "
         "the owner, creating additional delay."),
        ("Pros do not want AI talking directly to their customers. ",
         "This was a surprising and critical insight. My family was clear: as customers themselves, "
         "they don't like interacting with AI. They value the personal relationship with their service "
         "provider. \"I want to know I'm talking to a person, not a robot.\" This fundamentally shaped "
         "my approach \u2014 AI assists the pro behind the scenes; it never speaks to the customer directly."),
    ])

    doc.add_heading("Validating with Industry Data", level=2)
    para(doc,
        "After grounding the problem in firsthand research, I used AI tools (Claude, ChatGPT) to "
        "investigate whether these patterns hold across other home service verticals \u2014 plumbing, "
        "HVAC, electrical, cleaning, and general contracting. The data strongly confirmed my family's "
        "experience:"
    )
    bullets(doc, [
        ("60% of small business owners avoid confronting customers ",
         "about delinquent bills for fear of damaging the relationship (Old National Bank, 2025)."),
        ("56% of small businesses are owed money ",
         "from unpaid invoices, averaging $17,500 per business (QuickBooks 2025 Late Payments Report)."),
        ("Collection probability drops sharply with time: ",
         "94% at 30 days past due, 74% at 90 days, and only 50% at 6 months (Commercial Collection "
         "Agency Association). Every day of silence is measurably costly."),
        ("One reminder increases payment likelihood by ~50% ",
         "(Xero). AI-generated reminders get paid 5 days faster (QuickBooks 2025)."),
        ("SMS-delivered invoices are paid 32\u201340% faster ",
         "than email-only (Podium/Broadly data), yet HCP's automated reminders are email-only."),
        ("Dispute over work scope or pricing is the #1 reason ",
         "for non-payment in home services (SCORE/SBA)."),
    ])

    doc.add_heading("The Invoice Payment Funnel", level=2)
    para(doc,
        "Think of the invoice lifecycle as a funnel:"
    )
    para(doc,
        "Invoice Sent \u2192 Customer Receives \u2192 Questions / Concerns Arise \u2192 "
        "Pro Responds \u2192 Payment Made"
    )
    para(doc,
        "Today, the funnel leaks at the \"Questions / Concerns\" stage. Customers delay paying until "
        "their questions are answered. Pros delay answering because they're busy. The longer the gap, "
        "the lower the collection probability. AI can plug this leak \u2014 not by replacing the pro, "
        "but by having a response ready the moment they check their inbox."
    )

    doc.add_heading("What \"Good\" Looks Like", level=2)
    bullets(doc, [
        "Pro opens their inbox and immediately sees the 2\u20133 most important payment-related actions \u2014 no hunting.",
        "Each action has an AI-drafted message ready to send, grounded in actual conversation context and invoice details.",
        "Pro can see WHY the AI recommended this action (source messages, reasoning chain).",
        "Pro can send as-is or edit the draft \u2014 they stay in full control.",
        "Customers get timely, helpful responses about their invoices within hours, not days.",
        "Average days-to-payment decreases. The aging invoice tail shortens.",
        "Pro never has to make an uncomfortable \"where's my money?\" phone call.",
    ])

    doc.add_heading("What Failure Looks Like", level=2)
    bullets(doc, [
        "AI surfaces irrelevant or low-confidence actions \u2014 noise instead of signal.",
        "Drafted messages are inaccurate (wrong amounts, wrong customer) or tone-deaf.",
        "Pro loses trust and ignores the AI recommendations entirely.",
        "The feature overlaps with existing HCP automated reminders, creating customer confusion.",
        "Customers receive responses that feel robotic or impersonal.",
        "Pro feels they've lost control of their customer relationships.",
        "Dismiss rates are high, indicating recommendations aren't valuable to pros.",
    ])

    doc.add_page_break()

    # ── 2. SOLUTION DESIGN ──────────────────────────────────────────────────
    doc.add_heading("2. Solution Design", level=1)

    doc.add_heading("Approach: AI Smart Inbox", level=2)
    para(doc,
        "Enhance the existing HouseCall Pro messaging inbox with an intelligent layer that surfaces "
        "high-priority, payment-related actions as cards at the top of the inbox. Each card represents "
        "an AI-identified opportunity to move an invoice toward payment."
    )
    para(doc,
        "The core insight: HouseCall Pro already has incredibly rich context \u2014 conversation history, "
        "invoice details, job records, customer profiles. Today, the pro has to mentally synthesize all "
        "of this to decide what to do and what to say. AI can do this synthesis instantly and present "
        "the pro with a ready-to-send response, saving time and ensuring no opportunity is missed."
    )

    doc.add_heading("How It Works", level=2)

    bold_para(doc, "Action Cards (Top of Inbox)")
    para(doc,
        "Horizontally scrollable cards appear above the conversation list. Each card represents a "
        "specific, high-confidence AI recommendation based on real conversation analysis:"
    )
    numbered(doc, [
        ("Payment Promise Follow-ups \u2014 ",
         "Customer said \"I'll pay by Friday\" and Friday has passed. AI drafts a friendly, specific "
         "follow-up referencing the customer's own words about when they'd pay. Supports granular "
         "timing: \"end of day,\" \"this Friday,\" \"next week.\""),
        ("Invoice Question Responses \u2014 ",
         "Customer asked about a specific charge (e.g., \"What's the float charge?\"). AI drafts an "
         "explanation using the actual invoice line-item details and job context. Only surfaces when "
         "AI confidence is high enough to give an accurate answer."),
        ("Invoice Summaries \u2014 ",
         "When a new invoice is sent, AI drafts a customer-friendly breakdown highlighting key line "
         "items, the total, and how to pay \u2014 warmer and more informative than the raw invoice "
         "notification."),
        ("Payment Plan Suggestions \u2014 ",
         "Customer expressed financial difficulty (\"I can't pay this all at once\"). AI drafts a "
         "response acknowledging their concern and presenting available options like installment "
         "plans or financing through existing HCP integrations (e.g., Wisetack)."),
    ])

    bold_para(doc, "In-Conversation AI Drafts")
    para(doc,
        "When a pro opens a conversation that has an AI recommendation, the draft appears at the "
        "bottom of the chat thread \u2014 ready to review, edit, and send. The draft is presented "
        "in context alongside the actual conversation, so the pro can verify it makes sense."
    )

    bold_para(doc, "AI Reasoning Transparency")
    para(doc,
        "Every action card and in-conversation draft has a \"See reasoning\" button that opens a "
        "panel showing: the specific messages the AI used as context (highlighted with timestamps), "
        "the AI's reasoning chain explaining why this action was recommended, and the confidence "
        "level of the recommendation. This builds trust through transparency \u2014 the pro can "
        "verify the AI \"did its homework\" before sending anything."
    )

    doc.add_heading("What AI Does vs. What the Human Does", level=2)
    table(doc,
        ["AI Does", "Human Does"],
        [
            ["Scans conversations for payment-related signals", "Decides whether to send, edit, or dismiss"],
            ["Identifies follow-up timing based on customer promises", "Makes judgment calls on complex disputes"],
            ["Drafts contextual, personalized responses", "Approves every outbound message"],
            ["Surfaces invoice details relevant to customer questions", "Negotiates payment plans and custom terms"],
            ["Ranks actions by priority and confidence", "Owns the customer relationship"],
        ]
    )
    mixed(doc, [
        ("Key principle: ", True),
        ("AI is the drafting assistant. The Pro is the decision-maker. No message is ever sent "
         "without explicit Pro approval. This is non-negotiable in Phase 1 \u2014 a direct result "
         "of customer feedback that pros want to maintain personal relationships with their customers.", False),
    ])

    doc.add_heading("Non-AI Alternative Considered", level=2)
    bold_para(doc, "Rule-Based Follow-up System")
    para(doc, "I considered a deterministic, rule-based approach:")
    bullets(doc, [
        "Regex patterns to detect payment promises (\"pay\" + date keywords)",
        "Template responses for common invoice questions",
        "Scheduled reminders based on invoice age",
    ])

    bold_para(doc, "Why AI is the better choice:")
    bullets(doc, [
        ("Unstructured language. ",
         "Customer messages are natural language. \"I'll get you that money Friday,\" \"paying end "
         "of day,\" and \"can I settle up next week?\" all mean the same thing but require natural "
         "language understanding to parse reliably."),
        ("Context-dependent drafting. ",
         "Writing a good response requires understanding the specific charge being questioned, "
         "the customer's tone, the invoice details, and similar past interactions \u2014 not just "
         "pattern matching."),
        ("Semantic matching. ",
         "Connecting a customer's question \"What's the float charge?\" to the correct invoice "
         "line item (\"Floor leveling/float \u2014 $450\") is inherently a semantic task."),
        ("Scalability. ",
         "A rule-based system would need hundreds of hand-crafted rules and still miss edge cases. "
         "AI generalizes from context across any home service vertical."),
    ])
    para(doc,
        "That said, rule-based triggers can complement AI \u2014 for example, invoice age thresholds "
        "can signal when AI should look for follow-up opportunities, combining deterministic "
        "reliability with AI flexibility."
    )

    doc.add_page_break()

    # ── 3. TRUST AND FAILURE ────────────────────────────────────────────────
    doc.add_heading("3. Trust and Failure", level=1)

    doc.add_heading("Designing for When the System Is Wrong", level=2)
    para(doc,
        "The most important design decision in this product: when AI isn't confident, it stays "
        "silent. Surfacing a bad recommendation erodes trust faster than surfacing nothing. "
        "Below-threshold actions are never shown to the pro."
    )
    para(doc,
        "This principle came directly from customer research. My family told me they'd stop using "
        "a feature the moment it embarrassed them with a customer \u2014 one wrong amount, one "
        "tone-deaf follow-up to a sensitive situation, and the feature is dead to them. The trust "
        "bar in home services is extremely high because these are real, ongoing relationships built "
        "on years of service, not one-time transactions."
    )

    doc.add_heading("Guardrails for Every Interaction", level=2)
    para(doc, "Every AI draft is fully reviewable before it reaches a customer:")
    numbered(doc, [
        ("Full visibility. ", "The drafted message text is displayed in an editable text area. "
         "Nothing is hidden or sent automatically."),
        ("Reasoning transparency. ", "A \"See reasoning\" button shows the AI's logic and the "
         "specific source messages it analyzed \u2014 the pro can verify the AI understood the "
         "situation correctly."),
        ("Edit always available. ", "The pro can modify any part of the draft before sending. "
         "The AI's suggestion is a starting point, not a final answer."),
        ("Dismiss with signal. ", "A dismiss button removes the recommendation. Critically, "
         "dismissals are tracked as a quality signal (see Measurement section)."),
    ])

    doc.add_heading("Factual Grounding", level=2)
    para(doc,
        "AI drafts reference specific invoice line items, amounts, and dates from structured data "
        "\u2014 they are never generated from the model's training data alone. If the AI references "
        "\"$450 for floor leveling,\" the reasoning panel shows exactly which invoice line item it "
        "came from. Factual accuracy on financial details is a hard requirement: 100%, no exceptions. "
        "A single wrong dollar amount in a customer message could damage a pro's credibility and "
        "create a billing dispute."
    )

    doc.add_heading("How Trust Builds Over Time", level=2)
    para(doc,
        "The product is designed with a phased trust model. Each phase unlocks more AI autonomy, "
        "but only after the previous phase has proven the AI's reliability through measurable results:"
    )
    table(doc,
        ["Phase", "AI Behavior", "Pro Experience"],
        [
            ["Phase 1 (MVP)", "All drafts require manual review and explicit send",
             "Pro learns what AI can do, builds confidence through daily use"],
            ["Phase 2 (Earned)", "Pro can enable auto-send for specific, proven action types "
             "(e.g., invoice summaries)", "Earned autonomy \u2014 AI proved itself on simpler tasks first"],
            ["Phase 3 (Autonomous)", "Fully autonomous follow-ups for routine cases, with pro "
             "notification after the fact", "AI operates independently where trust is established"],
        ]
    )
    para(doc,
        "This MVP is Phase 1 only. Progression to Phase 2 is earned, not scheduled \u2014 it "
        "requires meeting quality bars, receiving positive qualitative feedback from pros, and "
        "demonstrating sustained accuracy over time. We do not graduate based on a calendar date; "
        "we graduate based on metrics."
    )

    doc.add_heading("Dismiss Tracking as a Trust Signal", level=2)
    para(doc,
        "Dismissal behavior is one of the strongest signals of whether AI recommendations are "
        "providing value. The design includes two levels of dismissal that give us granular insight:"
    )
    bullets(doc, [
        ("Individual recommendation dismiss \u2014 ",
         "\"This specific suggestion isn't helpful.\" Tells us the AI misjudged this particular "
         "situation. Acceptable at low rates."),
        ("Section dismiss (all action cards) \u2014 ",
         "\"I don't want AI suggestions at all.\" This is a critical signal: it means the overall "
         "feature isn't providing value to this pro, and we need to investigate why."),
    ])
    para(doc,
        "If per-recommendation dismiss rate exceeds 20%, we tighten the confidence threshold. "
        "If section-level dismissals exceed 5%, we pause rollout and conduct qualitative user "
        "research to understand the root cause before proceeding."
    )

    doc.add_page_break()

    # ── 4. MEASUREMENT ──────────────────────────────────────────────────────
    doc.add_heading("4. Measurement", level=1)

    doc.add_heading("Outcome Metrics", level=2)
    para(doc, "These measure whether the product moves the business needle:")
    table(doc,
        ["Metric", "Definition", "Target"],
        [
            ["Days-to-payment", "Average days from invoice sent to payment received",
             "Decrease by 20%+"],
            ["Response rate", "% of customer payment messages getting a pro response within 24 hours",
             "Increase from ~40% to 80%+"],
            ["30-day collection rate", "% of invoices paid within 30 days of sending",
             "Increase by 15%+"],
            ["Revenue recovered", "Dollar amount of invoices paid after AI-assisted follow-up",
             "Track absolute $"],
            ["Pro time saved", "Minutes per week pro spends on payment-related messaging",
             "Decrease by 50%+"],
        ]
    )

    doc.add_heading("AI Quality Metrics", level=2)
    para(doc, "These measure whether the AI itself is doing a good job:")
    table(doc,
        ["Metric", "Definition", "Quality Bar"],
        [
            ["Draft send rate", "% of AI drafts sent (as-is or edited) vs. dismissed", "> 70%"],
            ["Draft acceptance rate", "% of AI drafts sent without any editing", "> 50%"],
            ["Edit rate", "% of drafts edited before sending", "20\u201340% (healthy range)"],
            ["Dismissal rate", "% of action cards dismissed without sending", "< 20%"],
            ["Reasoning view rate", "% of actions where pro clicks \"See reasoning\"",
             "Track (expect higher early, declining as trust builds)"],
            ["Factual accuracy", "% of drafts with correct invoice amounts, dates, and line items",
             "100% (hard requirement)"],
        ]
    )

    doc.add_heading("Guardrail Metrics", level=2)
    para(doc,
        "These are protective thresholds. If any guardrail is breached, we take immediate action "
        "rather than waiting for the next review cycle:"
    )
    table(doc,
        ["Metric", "Threshold", "Action if Breached"],
        [
            ["Per-recommendation dismiss rate", "> 20%",
             "Investigate recommendation quality; tighten confidence threshold"],
            ["Section dismiss rate", "> 5%",
             "Pause rollout; conduct qualitative user research"],
            ["Customer satisfaction (CSAT)", "Drop > 5% vs. control",
             "Pause feature; investigate conversation quality"],
            ["Factual error rate", "Any occurrence",
             "Immediate hotfix; root cause analysis; incident review"],
            ["Pro NPS on AI feature", "< 30",
             "Re-evaluate approach and value proposition before expanding"],
        ]
    )

    doc.add_heading("Quality Bar Before Shipping", level=2)
    para(doc,
        "We do not proceed to wider rollout until all of the following are met in the A/B experiment:"
    )
    bullets(doc, [
        "Draft send rate > 70% \u2014 most drafts are relevant enough to act on",
        "Dismissal rate < 20% \u2014 the AI isn't spamming irrelevant actions",
        "Zero factual errors over a 30-day window \u2014 financial accuracy is non-negotiable",
        "Pro NPS on AI feature > 30 \u2014 pros find it genuinely helpful, not just tolerable",
        "No customer complaints about AI-generated message quality or tone",
        "Statistically significant improvement in days-to-payment in treatment vs. control group",
    ])

    doc.add_page_break()

    # ── 5. REQUIREMENTS SPECIFICATION ────────────────────────────────────────
    doc.add_heading("5. Requirements Specification", level=1)

    doc.add_heading("Product Vision", level=2)
    para(doc,
        "Enable HouseCall Pro's home service professionals to respond faster and more effectively "
        "to payment-related customer conversations using AI-powered draft suggestions \u2014 "
        "reducing days-to-payment and improving collection rates while keeping pros in full "
        "control of their customer relationships."
    )

    doc.add_heading("Goals", level=2)
    numbered(doc, [
        "Reduce average days-to-payment by 20% or more.",
        "Increase pro response rate on payment-related messages from ~40% to 80%+.",
        "Save pros 50%+ of time spent on payment-related messaging.",
        "Achieve >70% draft send rate, proving AI recommendations are valuable.",
        "Maintain 100% factual accuracy in every AI-generated draft.",
    ])

    doc.add_heading("Non-Goals", level=2)
    bullets(doc, [
        "Replacing the pro in customer communication. Phase 1 is human-in-the-loop only; AI never sends messages autonomously.",
        "Changing payment infrastructure. HCP's payment rails (cards, ACH, financing, BNPL, tap-to-pay) are not the problem.",
        "Replacing HCP's existing automated email reminders. This feature complements reminders, not duplicates them.",
        "Supporting non-payment conversations (scheduling, estimates, general inquiries). Out of scope for MVP.",
        "Multi-user team collaboration or role-based access. Out of scope for MVP.",
        "Push notifications or real-time alerting. Out of scope for MVP.",
    ])

    doc.add_heading("Target Personas", level=2)

    doc.add_heading("Persona 1: \"Mike\" \u2014 The Solo Plumber", level=3)
    table(doc,
        ["Attribute", "Detail"],
        [
            ["Profile", "38, one-man plumbing business, 4\u20136 jobs/day, ~$180K/year revenue"],
            ["Pain", "~$12K in unpaid invoices. Sends one manual follow-up, then drops it. "
             "Hates feeling like a bill collector."],
            ["Message overload", "15\u201320 customer texts/day through HCP. Payment messages "
             "get lost in scheduling chatter."],
        ]
    )

    doc.add_heading("Persona 2: \"Maria\" \u2014 The Growing Cleaning Company Owner", level=3)
    table(doc,
        ["Attribute", "Detail"],
        [
            ["Profile", "42, 12-employee cleaning company, ~$450K/year revenue"],
            ["Pain", "15\u201325% of invoices overdue at any time (~$8K\u2013$15K). Office helper "
             "can't answer billing questions \u2014 escalates to Maria, who's busy."],
            ["Message overload", "Multiple team members see messages but nobody \"owns\" follow-up. "
             "Questions sit unanswered for days."],
        ]
    )

    doc.add_heading("Persona 3: \"Jesse\" \u2014 The HVAC Technician-Owner", level=3)
    table(doc,
        ["Attribute", "Detail"],
        [
            ["Profile", "51, 4-person HVAC company, ~$600K/year revenue, high-ticket jobs ($2K\u2013$8K)"],
            ["Pain", "$35K in outstanding receivables. Only Jesse can explain technical pricing "
             "differences \u2014 he's the bottleneck."],
            ["Message overload", "Critical questions like \"why is this $800 more than the estimate?\" "
             "sit for 2\u20133 days while Jesse is on a roof."],
        ]
    )

    doc.add_heading("Functional Requirements", level=2)

    doc.add_heading("P0 \u2014 Must Have (MVP)", level=3)
    table(doc,
        ["ID", "Requirement", "Rationale"],
        [
            ["FR-01", "AI scans conversations to identify payment-related signals "
             "(promises, questions, concerns, new invoices)",
             "Core value proposition \u2014 surface what matters without pro effort"],
            ["FR-02", "Action cards displayed at top of inbox for high-confidence "
             "recommendations with priority ranking",
             "Immediate visibility; pro sees actions without searching"],
            ["FR-03", "AI generates contextual draft responses using conversation "
             "history and structured invoice data",
             "Core feature \u2014 ready-to-send messages grounded in real data"],
            ["FR-04", "Pro can review, edit, and send AI-drafted messages",
             "Human-in-the-loop control; no autonomous sending in Phase 1"],
            ["FR-05", "Reasoning transparency: source messages, confidence score, "
             "and logic visible via \"See reasoning\" panel",
             "Trust-building through transparency; pro verifies AI \"did its homework\""],
            ["FR-06", "Dismiss functionality for individual action cards",
             "Pro control; dismissals tracked as quality signal"],
            ["FR-07", "Factual grounding: all amounts, dates, and line items "
             "sourced from structured invoice data, never hallucinated",
             "Financial accuracy is non-negotiable \u2014 zero tolerance for errors"],
        ]
    )

    doc.add_heading("P1 \u2014 Should Have", level=3)
    table(doc,
        ["ID", "Requirement", "Rationale"],
        [
            ["FR-08", "Four distinct action types: payment follow-up, invoice question, "
             "invoice summary, payment plan",
             "Different situations need different response strategies and tone"],
            ["FR-09", "Priority ranking of action cards (high / medium / low) with "
             "visual differentiation",
             "Help pros triage \u2014 highest-impact actions are most visible"],
            ["FR-10", "Section-level dismiss (opt out of all AI suggestions)",
             "Respect pro autonomy; strong negative signal for guardrail monitoring"],
            ["FR-11", "Dismiss event tracking for analytics and guardrail metrics",
             "Quantitative signal for whether recommendations provide value"],
        ]
    )

    doc.add_heading("P2 \u2014 Nice to Have (Post-MVP)", level=3)
    table(doc,
        ["ID", "Requirement", "Rationale"],
        [
            ["FR-12", "Thumbs up / down feedback on individual drafts",
             "Explicit quality signal to complement implicit send/dismiss data"],
            ["FR-13", "Draft improvement learning from pro edit patterns",
             "Continuous quality improvement based on how pros modify AI suggestions"],
            ["FR-14", "Customizable confidence threshold per pro",
             "Power users may want more or fewer suggestions"],
            ["FR-15", "Auto-send for specific action types with pro opt-in (Phase 2)",
             "Earned autonomy for proven, low-risk action types"],
        ]
    )

    doc.add_heading("Non-Functional Requirements", level=2)
    table(doc,
        ["ID", "Requirement", "Target"],
        [
            ["NFR-01", "Draft generation latency", "< 5 seconds from request to display"],
            ["NFR-02", "Action card refresh frequency", "Within 1 minute of new incoming messages"],
            ["NFR-03", "Platform support", "Mobile-first (iOS and Android); responsive web"],
            ["NFR-04", "Accessibility", "WCAG 2.1 AA compliance"],
            ["NFR-05", "Data privacy", "No customer PII used in AI model training; all processing "
             "via API (no fine-tuning on customer data)"],
            ["NFR-06", "Availability", "99.9% uptime; graceful degradation if AI service is unavailable "
             "(inbox still functional without AI features)"],
        ]
    )

    doc.add_heading("Key Metrics", level=2)
    doc.add_heading("Primary Success Metrics", level=3)
    table(doc,
        ["Metric", "Definition", "Target"],
        [
            ["Days-to-payment", "Avg days from invoice sent to payment received", "Decrease by 20%+"],
            ["Pro response rate", "% of payment messages responded to within 24h", "40% \u2192 80%+"],
            ["30-day collection rate", "% of invoices paid within 30 days", "Increase by 15%+"],
            ["Draft send rate", "% of AI drafts sent (as-is or edited)", "> 70%"],
            ["Factual accuracy", "% of drafts with correct financial details", "100%"],
        ]
    )

    doc.add_heading("Guardrail Metrics", level=3)
    table(doc,
        ["Metric", "Red Line", "Response"],
        [
            ["Per-card dismiss rate", "> 20%", "Tighten confidence threshold; review prompt quality"],
            ["Section dismiss rate", "> 5%", "Pause rollout; qualitative research"],
            ["Customer CSAT delta", "Drop > 5% vs. control", "Pause feature; investigate"],
            ["Factual errors", "Any occurrence", "Immediate hotfix; incident review"],
            ["Pro NPS on feature", "< 30", "Re-evaluate before expanding"],
        ]
    )

    doc.add_heading("Rollout Strategy", level=2)

    doc.add_heading("Phase 0: Internal Dogfooding (2 weeks)", level=3)
    bullets(doc, [
        "HCP product and engineering team tests with synthetic conversations.",
        "Identify obvious UX issues, edge cases, and prompt failures before any customer exposure.",
        "Build internal confidence and gather team feedback on draft quality.",
    ])

    doc.add_heading("Phase 1a: Closed Beta (4 weeks)", level=3)
    bullets(doc, [
        "50\u2013100 hand-selected pros across verticals (plumbing, HVAC, electrical, cleaning, flooring).",
        "Mix of solo operators and small teams to test both persona types.",
        "Weekly 15-minute feedback sessions with participating pros.",
        "Quality gate to proceed: draft send rate > 60%, dismiss rate < 25%, zero factual errors.",
    ])

    doc.add_heading("Phase 1b: A/B Experiment (6\u20138 weeks)", level=3)
    bullets(doc, [
        "~1,000 pros randomly assigned to treatment (AI Smart Inbox) vs. control (standard inbox).",
        "Measure days-to-payment, response rate, collection rate, and pro satisfaction.",
        "Statistical significance required on primary metrics before proceeding.",
        "Quality gate to proceed: draft send rate > 70%, dismiss rate < 20%, zero factual errors, "
        "NPS > 30, statistically significant improvement in days-to-payment.",
    ])

    doc.add_heading("Phase 1c: General Availability", level=3)
    bullets(doc, [
        "Gradual rollout to all HCP pros (10% \u2192 25% \u2192 50% \u2192 100%).",
        "Feature flag for instant kill switch if guardrails are breached.",
        "Continuous monitoring of all guardrail metrics at each rollout stage.",
        "Qualitative feedback collection through in-app surveys at 7 and 30 days post-enablement.",
    ])

    doc.add_heading("Phase 2: Earned Autonomy (Future)", level=3)
    para(doc,
        "Phase 2 is not on a fixed timeline. It is earned by sustained performance:"
    )
    bullets(doc, [
        "Prerequisites: 90+ days of sustained quality metrics at scale, explicit pro demand for "
        "more automation, zero customer complaints about AI quality.",
        "Scope: enable opt-in auto-send for specific, low-risk action types (e.g., invoice summaries).",
        "Each action type graduates independently based on its own track record.",
        "Pros can opt in or out at any time; autonomy is always revocable.",
    ])

    doc.add_heading("Phase Transition Criteria", level=3)
    table(doc,
        ["Gate", "Criteria to Pass"],
        [
            ["Internal \u2192 Beta", "No critical bugs; team consensus on draft quality"],
            ["Beta \u2192 A/B", "Draft send rate > 60%; dismiss rate < 25%; zero factual errors; "
             "qualitative feedback positive (NPS > 20)"],
            ["A/B \u2192 GA", "Statistically significant improvement in days-to-payment; draft "
             "send rate > 70%; dismiss rate < 20%; NPS > 30; zero factual errors"],
            ["Phase 1 \u2192 Phase 2", "90+ days sustained quality at scale; explicit pro demand "
             "for automation; zero customer complaints"],
        ]
    )

    doc.save(os.path.join(OUTPUT_DIR, "AI_Smart_Inbox_Process_Document.docx"))
    print("  [1/3] Process Document created.")


# ── Document 3: AI Evaluation Spec ──────────────────────────────────────────

def create_eval_spec():
    doc = new_doc()
    title_page(doc, "AI Smart Inbox\nfor Invoice Collection", "AI Evaluation Specification")

    # ── Purpose ──────────────────────────────────────────────────────────────
    doc.add_heading("Purpose and Scope", level=1)
    para(doc,
        "This document defines how we evaluate the AI system powering the Smart Inbox \u2014 "
        "covering draft quality, factual accuracy, relevance, and tone. The eval spec serves "
        "two purposes: (1) establishing quality gates that must be met before any customer "
        "exposure, and (2) defining the ongoing monitoring framework that ensures AI quality "
        "is sustained as the product scales."
    )
    para(doc,
        "The fundamental question we need to answer at every stage: are the AI's recommendations "
        "genuinely helping pros get paid faster, or are they creating noise?"
    )

    # ── What We're Evaluating ────────────────────────────────────────────────
    doc.add_heading("What We Are Evaluating", level=1)
    para(doc, "The AI system performs four functions that each require distinct evaluation:")
    table(doc,
        ["Function", "What It Does", "Key Risk"],
        [
            ["Action Detection", "Identifies which conversations need a payment-related action",
             "False positives (surfacing irrelevant conversations) or false negatives (missing real opportunities)"],
            ["Action Classification", "Categorizes the action type (follow-up, question, summary, payment plan)",
             "Misclassification leads to wrong draft tone and content"],
            ["Draft Generation", "Writes a contextual response for the pro to review",
             "Inaccurate facts, wrong tone, unhelpful content, or hallucinated details"],
            ["Confidence Scoring", "Assigns a confidence level to each recommendation",
             "Poor calibration \u2014 high-confidence scores on bad drafts, or overly conservative filtering"],
        ]
    )

    # ── Evaluation Dimensions ────────────────────────────────────────────────
    doc.add_heading("Evaluation Dimensions", level=1)
    para(doc,
        "Every AI draft is evaluated across five dimensions. Each dimension has a clear definition "
        "and scoring rubric to ensure consistency across evaluators:"
    )

    doc.add_heading("1. Factual Accuracy (Pass / Fail)", level=2)
    para(doc,
        "Are all financial details \u2014 invoice amounts, line items, dates, customer names \u2014 "
        "correct and sourced from structured data? This is a binary pass/fail dimension. A single "
        "factual error fails the entire draft regardless of other dimension scores."
    )
    bullets(doc, [
        "Invoice total matches the actual invoice",
        "Line item descriptions and amounts are correct",
        "Dates (due dates, promise dates, invoice dates) are accurate",
        "Customer name is correct",
        "No hallucinated charges, services, or details",
    ])

    doc.add_heading("2. Relevance (1\u20135 scale)", level=2)
    table(doc,
        ["Score", "Definition"],
        [
            ["5", "Draft directly addresses the customer's specific question or situation with precision"],
            ["4", "Draft addresses the right topic but could be more specific to this customer's context"],
            ["3", "Draft is related but generic \u2014 could apply to many customers, not tailored"],
            ["2", "Draft partially addresses the situation but misses the main point"],
            ["1", "Draft is irrelevant to the conversation or addresses the wrong issue entirely"],
        ]
    )

    doc.add_heading("3. Tone (1\u20135 scale)", level=2)
    table(doc,
        ["Score", "Definition"],
        [
            ["5", "Warm, professional, natural \u2014 reads like it came from the pro, not a bot"],
            ["4", "Appropriate tone but slightly generic; could be warmer or more personalized"],
            ["3", "Neutral/acceptable but noticeably \"template-like\""],
            ["2", "Tone mismatch (too formal, too casual, or insensitive to the situation)"],
            ["1", "Tone is inappropriate (aggressive, dismissive, or robotic for the context)"],
        ]
    )

    doc.add_heading("4. Helpfulness (1\u20135 scale)", level=2)
    table(doc,
        ["Score", "Definition"],
        [
            ["5", "Sending this draft as-is would likely move the invoice toward payment"],
            ["4", "Helpful with minor edits needed; addresses the customer's needs"],
            ["3", "Somewhat helpful but requires significant editing to be sendable"],
            ["2", "Minimally helpful; pro would be better off writing from scratch"],
            ["1", "Unhelpful or counterproductive; could damage the customer relationship"],
        ]
    )

    doc.add_heading("5. Completeness (1\u20135 scale)", level=2)
    table(doc,
        ["Score", "Definition"],
        [
            ["5", "Addresses all customer concerns; includes relevant invoice context and next steps"],
            ["4", "Addresses main concern; minor points not covered but not critical"],
            ["3", "Addresses some concerns but misses important context or follow-up"],
            ["2", "Incomplete \u2014 addresses only part of what the customer asked"],
            ["1", "Fails to address the customer's actual concern"],
        ]
    )

    doc.add_page_break()

    # ── Offline Evaluation ───────────────────────────────────────────────────
    doc.add_heading("Offline Evaluation (Pre-Launch)", level=1)
    para(doc,
        "Before any customer-facing deployment, the AI system is tested against a curated "
        "test set that covers the full range of expected inputs, including edge cases."
    )

    doc.add_heading("Test Set Construction", level=2)
    bullets(doc, [
        ("50+ synthetic conversation scenarios ", "spanning all 4 action types, built from "
         "real-world patterns observed in customer research."),
        ("Edge cases included: ", "ambiguous customer intent, multiple questions in one message, "
         "emotional or frustrated language, incorrect customer assumptions about charges, "
         "conversations with insufficient context for a recommendation."),
        ("Gold standard responses ", "written by domain experts (real home service pros or "
         "experienced HCP support staff) for each scenario."),
        ("Negative examples: ", "conversations where the AI should not generate a recommendation "
         "(e.g., paid invoices, scheduling-only threads) to test false positive rates."),
    ])

    doc.add_heading("Evaluation Process", level=2)
    numbered(doc, [
        "Generate AI drafts for each scenario in the test set.",
        "Three independent human annotators rate each draft on all five dimensions.",
        "Measure inter-annotator agreement using Cohen's kappa (\u03ba).",
        "Factual accuracy checked programmatically against structured invoice data.",
        "Compare AI drafts to gold standard responses on relevance and completeness.",
    ])

    doc.add_heading("Pass Criteria for Launch", level=2)
    table(doc,
        ["Dimension", "Threshold"],
        [
            ["Factual accuracy", "100% pass rate (zero failures)"],
            ["Average relevance", "\u2265 4.0 / 5"],
            ["Average tone", "\u2265 4.0 / 5"],
            ["Average helpfulness", "\u2265 3.5 / 5"],
            ["Average completeness", "\u2265 3.5 / 5"],
            ["Inter-annotator agreement", "\u03ba \u2265 0.7"],
            ["False positive rate (action detection)", "< 5%"],
            ["Action type classification accuracy", "\u2265 95%"],
        ]
    )

    # ── Online Evaluation ────────────────────────────────────────────────────
    doc.add_heading("Online Evaluation (Post-Launch)", level=1)
    para(doc,
        "Once the AI is live, we shift to continuous monitoring using both implicit behavioral "
        "signals and explicit pro feedback."
    )

    doc.add_heading("Implicit Behavioral Signals", level=2)
    table(doc,
        ["Signal", "What It Tells Us", "How We Measure"],
        [
            ["Send rate", "Overall draft quality and relevance",
             "% of drafts sent (as-is or edited) vs. dismissed"],
            ["Edit rate", "How close drafts are to \"ready to send\"",
             "% of drafts edited before sending"],
            ["Edit distance", "How much pros change the draft content",
             "Levenshtein distance between AI draft and sent message"],
            ["Dismiss rate", "Irrelevance or quality problems",
             "% of action cards dismissed without action"],
            ["Time-to-action", "Whether AI speeds up pro response",
             "Time from action card appearance to pro response vs. baseline"],
            ["Payment outcome", "Whether AI-assisted conversations get paid faster",
             "Days-to-payment for AI-assisted vs. non-assisted conversations"],
        ]
    )

    doc.add_heading("Explicit Feedback Signals", level=2)
    bullets(doc, [
        ("Thumbs up / down on drafts (P2 feature). ",
         "Direct quality signal from the pro on individual recommendations."),
        ("Dismiss reason (optional). ",
         "When a pro dismisses, a lightweight prompt asks why: \"Not relevant,\" "
         "\"Wrong tone,\" \"Inaccurate info,\" or \"Already handled.\""),
        ("Periodic NPS surveys. ",
         "At 7 and 30 days post-enablement, ask: \"How likely are you to recommend "
         "the AI inbox to another pro?\""),
    ])

    # ── Confidence Calibration ───────────────────────────────────────────────
    doc.add_heading("Confidence Calibration", level=1)
    para(doc,
        "The AI assigns a confidence score (0\u2013100%) to each recommendation. For the system "
        "to be trustworthy, stated confidence must correlate with actual quality:"
    )
    bullets(doc, [
        "A draft with 90% confidence should be sent (as-is or edited) approximately 90% of the time.",
        "Track confidence vs. send rate weekly. If they diverge (e.g., 90% confidence but only 60% "
        "send rate), the confidence model needs recalibration.",
        "Adjust the minimum confidence threshold for showing recommendations based on observed "
        "performance. Start conservative (higher threshold) and lower it as the system proves itself.",
    ])

    # ── Prompt Regression Testing ────────────────────────────────────────────
    doc.add_heading("Prompt Regression Testing", level=1)
    para(doc,
        "Every change to the AI prompts (system prompts, user prompt templates, or model parameters) "
        "goes through regression testing before deployment:"
    )
    bullets(doc, [
        ("Fixed critical test set: ", "20 \"golden\" scenarios that represent the most important "
         "and most common situations. Any degradation on these blocks the prompt change."),
        ("Full test set run: ", "All 50+ scenarios re-evaluated after significant prompt changes. "
         "Results compared to baseline."),
        ("Prompt versioning: ", "Every prompt change is version-controlled with a changelog "
         "documenting what changed and why."),
        ("A/B testing for major changes: ", "Significant prompt rewrites are tested in a shadow "
         "mode (generate both old and new drafts, compare quality) before replacing the live prompt."),
    ])

    # ── Monitoring & Alerting ────────────────────────────────────────────────
    doc.add_heading("Monitoring and Alerting", level=1)

    doc.add_heading("Real-Time Dashboards", level=2)
    bullets(doc, [
        "Draft send rate, edit rate, and dismiss rate (rolling 24h, 7d, 30d)",
        "Confidence score distribution (are scores clustering? shifting?)",
        "Factual error reports (any pro-reported inaccuracies)",
        "Action type breakdown (which types are performing best/worst?)",
        "Latency (draft generation time p50, p95, p99)",
    ])

    doc.add_heading("Automated Alerts", level=2)
    table(doc,
        ["Condition", "Alert Level", "Response"],
        [
            ["Dismiss rate > 20% (rolling 24h)", "Warning",
             "Investigate; review recent drafts for quality"],
            ["Any factual error reported", "Critical",
             "Immediate investigation; potential feature pause"],
            ["Send rate drops > 10% week-over-week", "Warning",
             "Review prompt changes; check for model behavior shifts"],
            ["Draft generation latency p95 > 10s", "Warning",
             "Investigate API performance; consider timeout adjustments"],
            ["Section dismiss rate > 5%", "Critical",
             "Pause rollout expansion; schedule user research sessions"],
        ]
    )

    # ── Iteration Cadence ────────────────────────────────────────────────────
    doc.add_heading("Evaluation Cadence", level=1)
    table(doc,
        ["Cadence", "Activity"],
        [
            ["Daily", "Automated dashboard review; alert triage"],
            ["Weekly", "Manual review of 10\u201320 randomly sampled drafts across action types; "
             "confidence calibration check"],
            ["Monthly", "Full eval run against complete test set; analyze edit patterns for "
             "prompt improvement opportunities; publish quality report"],
            ["Quarterly", "Refresh test set with new real-world scenarios discovered in production; "
             "recalibrate confidence thresholds; review whether Phase 2 criteria are met"],
        ]
    )

    # ── Connection to Rollout ────────────────────────────────────────────────
    doc.add_heading("Eval Metrics That Gate Rollout Decisions", level=1)
    para(doc,
        "The eval spec is not academic \u2014 it directly controls product rollout decisions. "
        "The following eval outcomes gate each phase transition:"
    )
    table(doc,
        ["Transition", "Eval Requirements"],
        [
            ["Offline \u2192 Beta", "Pass all offline eval thresholds (100% factual accuracy, "
             "\u22654.0 relevance, \u22654.0 tone, \u22653.5 helpfulness)"],
            ["Beta \u2192 A/B", "Online send rate >60%, dismiss rate <25%, zero factual errors, "
             "positive qualitative feedback from beta pros"],
            ["A/B \u2192 GA", "Statistically significant improvement in days-to-payment, "
             "online send rate >70%, dismiss rate <20%, confidence calibration within 10% of actual"],
            ["GA \u2192 Phase 2", "90+ days of sustained metrics at scale, no factual errors, "
             "NPS > 30, explicit pro demand for auto-send"],
        ]
    )

    doc.save(os.path.join(OUTPUT_DIR, "AI_Evaluation_Spec.docx"))
    print("  [2/3] AI Evaluation Spec created.")


# ── Document 4: AI Usage Log ────────────────────────────────────────────────

def create_usage_log():
    doc = new_doc()
    title_page(doc, "AI Smart Inbox\nfor Invoice Collection", "AI Usage Log")

    doc.add_heading("Overview", level=1)
    para(doc,
        "This document records how AI tools were used throughout the development of the AI Smart "
        "Inbox prototype \u2014 from initial research through implementation and testing. AI was "
        "used as an accelerator and thought partner at every phase, but all strategic decisions, "
        "product direction, and quality judgments were made by the product manager."
    )

    doc.add_heading("Tools Used", level=2)
    table(doc,
        ["Tool", "Primary Use"],
        [
            ["Claude (Anthropic)", "Problem research, competitive analysis, requirements refinement, "
             "idea feedback, prompt engineering collaboration"],
            ["ChatGPT / OpenAI", "Problem research, feature understanding, customer feedback analysis, "
             "data validation"],
            ["Claude Code", "Primary code generation tool for the working prototype (frontend, "
             "backend, and AI integration)"],
            ["OpenAI API (GPT-5.4)", "In-product AI: generates contextual draft responses via "
             "structured output with Zod schemas"],
        ]
    )

    doc.add_page_break()

    # ── Phase 1 ──────────────────────────────────────────────────────────────
    doc.add_heading("Phase 1: Research and Discovery", level=1)

    doc.add_heading("What I Did", level=2)
    bullets(doc, [
        ("Primary user research (human-led): ",
         "Conducted in-depth interviews with my father and stepmother, who own a small flooring "
         "company (3 employees). Walked through their complete invoice-to-payment workflow, asking "
         "open-ended questions about pain points, workarounds, and what an ideal solution would "
         "look like. This revealed the core insight that customer questions about invoice line items "
         "are the primary bottleneck to payment."),
        ("Problem research with AI: ",
         "Used Claude and ChatGPT to investigate whether the patterns I observed in my family's "
         "business held across other home service verticals (plumbing, HVAC, electrical, cleaning). "
         "AI helped me find and synthesize industry data on late payments, collection probabilities, "
         "and SMS vs. email payment rates."),
        ("HouseCall Pro feature understanding: ",
         "Used AI to understand HCP's existing product ecosystem \u2014 what payment methods they "
         "support, how their messaging/inbox works, what automated reminders exist, and where gaps "
         "remain."),
        ("Competitive analysis: ",
         "Used AI to map the competitive landscape: how other field service platforms (Jobber, "
         "ServiceTitan, Kickserv) handle invoice follow-ups, and how fintech tools (QuickBooks, "
         "Xero, FreshBooks) approach AI-assisted collections."),
        ("Idea feedback: ",
         "Brainstormed multiple solution concepts with AI, getting feedback on feasibility, "
         "differentiation, and potential customer impact. Evaluated approaches ranging from "
         "fully automated follow-ups to the human-in-the-loop model I ultimately chose."),
    ])

    doc.add_heading("What Required Human Judgment", level=2)
    bullets(doc, [
        "Choosing to start with family interviews rather than secondary research \u2014 talking to real "
        "people first grounded the entire project in authentic customer experience.",
        "Interpreting my family's emotional pushback against AI-to-customer communication \u2014 AI "
        "tools didn't surface this insight; a conversation at the kitchen table did.",
        "Deciding which of several possible pain points to focus on (invoice questions and follow-ups "
        "vs. scheduling optimization, vs. estimate accuracy).",
    ])

    # ── Phase 2 ──────────────────────────────────────────────────────────────
    doc.add_heading("Phase 2: Requirements and Design", level=1)

    doc.add_heading("What I Did", level=2)
    bullets(doc, [
        ("User scenario definition (human-led, AI-enhanced): ",
         "I defined the core user scenarios and action types based on customer research. Used AI "
         "to stress-test the scenarios: \"What edge cases am I missing?\" \"What could go wrong?\" "
         "\"Is this the right granularity of action types?\""),
        ("Requirements prioritization: ",
         "I set the P0/P1/P2 priorities based on what I believed would deliver the most value with "
         "the least risk. AI helped me think through trade-offs and identify dependencies I might "
         "have missed."),
        ("UI and interaction design (human-led): ",
         "I defined the interaction model (action cards at top of inbox, in-conversation drafts, "
         "reasoning panel) and chose to present it as an iPhone mockup to match HCP's mobile-first "
         "user base. AI helped refine specific interaction details."),
        ("Technology stack selection (human decision): ",
         "I chose React + Vite for the frontend, Express for the backend, and OpenAI's API for "
         "draft generation based on my assessment of what would enable the fastest path to a "
         "working prototype."),
    ])

    doc.add_heading("What Required Human Judgment", level=2)
    bullets(doc, [
        "The decision that Phase 1 must be 100% human-in-the-loop (no auto-sending) was a direct "
        "result of my family's feedback. AI tools generally suggested a more aggressive automation "
        "approach.",
        "Choosing to include dismiss tracking as a guardrail metric \u2014 this came from thinking "
        "about how to measure trust, not from AI suggestion.",
        "Scoping the MVP to four specific action types rather than trying to cover all possible "
        "inbox scenarios.",
    ])

    # ── Phase 3 ──────────────────────────────────────────────────────────────
    doc.add_heading("Phase 3: Implementation", level=1)

    doc.add_heading("What I Did", level=2)
    bullets(doc, [
        ("Code generation with Claude Code: ",
         "Claude Code was the primary tool for building the working prototype. I described what I "
         "wanted (component structure, interaction flow, visual design) and Claude Code generated "
         "the code. This included the React frontend (inbox view, conversation view, AI reasoning "
         "modal, iPhone mockup frame), the Express backend (REST API, mock data), and the OpenAI "
         "integration (structured output with Zod schemas)."),
        ("Prompt engineering (collaborative): ",
         "I worked with AI to define the four specialized system prompts for each action type "
         "(payment follow-up, invoice question, invoice summary, payment plan). The process was "
         "iterative: I'd describe the desired tone and content, review the generated prompts, test "
         "them against mock scenarios, and refine based on output quality."),
        ("Mock data design (human-led): ",
         "I defined the six conversation scenarios in the mock data, ensuring they covered the "
         "range of action types and included one conversation (Rachel Thompson) with no AI action "
         "to demonstrate selectivity. Claude Code implemented the data structures."),
    ])

    doc.add_heading("What Required Human Judgment", level=2)
    bullets(doc, [
        "Deciding what mock scenarios to include and ensuring they told a coherent story about "
        "the product's capabilities.",
        "Evaluating whether the AI-generated draft responses felt natural and appropriate for "
        "a blue-collar, relationship-driven industry.",
        "Choosing the visual design direction (iOS-inspired, clean, professional) to align with "
        "HCP's existing mobile app aesthetic.",
    ])

    # ── Phase 4 ──────────────────────────────────────────────────────────────
    doc.add_heading("Phase 4: Testing and Iteration", level=1)

    doc.add_heading("What I Did", level=2)
    bullets(doc, [
        ("Manual testing (human-led): ",
         "I conducted all testing manually \u2014 walking through the complete user flow, checking "
         "each action type, testing edge cases, verifying that AI drafts were factually grounded "
         "in the mock invoice data."),
        ("Bug fixing with Claude Code: ",
         "When I identified issues during testing, I described them to Claude Code and used it to "
         "implement fixes. This included UI polish, interaction bugs, and edge cases in the AI "
         "draft display."),
        ("Prompt iteration: ",
         "Based on observing draft quality across different scenarios, I refined the system prompts "
         "to improve tone, specificity, and factual grounding."),
    ])

    doc.add_heading("What Required Human Judgment", level=2)
    bullets(doc, [
        "Evaluating draft quality is inherently subjective \u2014 does this sound like something a "
        "real pro would send? AI can't make this judgment for its own output.",
        "Deciding which issues were worth fixing in the MVP vs. noting for future work.",
        "Assessing whether the overall experience was compelling enough to demonstrate the product concept.",
    ])

    doc.add_page_break()

    # ── Reflection ───────────────────────────────────────────────────────────
    doc.add_heading("Reflection: How AI Changed My Process", level=1)

    doc.add_heading("Where AI Added the Most Value", level=2)
    bullets(doc, [
        ("Speed of research: ",
         "AI compressed weeks of competitive analysis and data gathering into days. I could "
         "quickly validate whether patterns I observed in one vertical held across others."),
        ("Rapid prototyping: ",
         "Claude Code enabled me to go from concept to working prototype dramatically faster "
         "than traditional development. The prototype is functional, not a mockup."),
        ("Thought partnership: ",
         "AI was a useful sounding board for stress-testing ideas. \"What am I missing?\" and "
         "\"What could go wrong?\" prompts surfaced considerations I might have overlooked."),
        ("Prompt engineering: ",
         "Collaborating with AI on writing the prompts for the product's own AI features was "
         "uniquely effective \u2014 AI helping design AI."),
    ])

    doc.add_heading("Where Human Judgment Was Irreplaceable", level=2)
    bullets(doc, [
        ("Customer empathy and intuition: ",
         "The critical insight that pros don't want AI talking to their customers came from a "
         "kitchen-table conversation, not a data analysis. Understanding the emotional weight "
         "of invoice follow-ups \u2014 the fear of damaging a relationship \u2014 required human "
         "empathy."),
        ("Problem selection: ",
         "Choosing which problem to solve, among many possible options, required product judgment "
         "about market need, technical feasibility, and alignment with HCP's strategic direction."),
        ("Design decisions: ",
         "The phased trust model, the dismiss-as-guardrail concept, and the \"silence is better "
         "than noise\" principle were all products of human strategic thinking."),
        ("Quality evaluation: ",
         "Judging whether an AI draft \"sounds right\" for a blue-collar pro communicating with "
         "a homeowner is a fundamentally human assessment."),
    ])

    doc.add_heading("Key Takeaway", level=2)
    para(doc,
        "AI is most valuable as an accelerator and thought partner, not a replacement for product "
        "thinking. The most important decisions in this project \u2014 the problem to solve, the "
        "approach to take, the guardrails to enforce \u2014 all required human judgment informed by "
        "real customer conversations. AI made those decisions faster and better-informed, but it "
        "didn't make them for me."
    )

    doc.save(os.path.join(OUTPUT_DIR, "AI_Usage_Log.docx"))
    print("  [3/3] AI Usage Log created.")


# ── Main ─────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    # Remove old separate requirements file if it exists
    old_req = os.path.join(OUTPUT_DIR, "Requirements_Specification.docx")
    if os.path.exists(old_req):
        os.remove(old_req)
        print(f"  Removed old separate file: Requirements_Specification.docx")
    print(f"Generating documents in {OUTPUT_DIR}...\n")
    create_process_document()
    create_eval_spec()
    create_usage_log()
    print(f"\nDone! All documents saved to: {OUTPUT_DIR}/")
